# utils/report_generator.py
import os
import datetime
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class PatientReportGenerator:
    """
    Utility class for generating patient reports in Word format.
    Includes patient information, test results, and visit history.
    """
    
    def __init__(self, config_manager):
        self.config_manager = config_manager
    
    def generate_report(self, patient_data, visit_history=None, test_results=None, include_doctor_notes=True):
        """
        Generate a patient report as a Word document.
        
        Args:
            patient_data (dict): Patient information
            visit_history (list): List of visit records
            test_results (list): List of test results
            include_doctor_notes (bool): Whether to include space for doctor's notes
            
        Returns:
            str: Path to the generated document
        """
        # Create a new Word document
        doc = Document()
        
        # Set up the document
        self._setup_document(doc)
        
        # Add clinic information
        self._add_clinic_info(doc)
        
        # Add report title
        self._add_title(doc, f"Patient Medical Report")
        
        # Add patient information section
        self._add_patient_info(doc, patient_data)
        
        # Add medications section if available
        medications = patient_data.get("medications", [])
        if medications:
            self._add_medications(doc, medications)
            
        # Add medical history if available
        medical_history = patient_data.get("medical_history", [])
        if medical_history:
            self._add_medical_history(doc, medical_history)
        
        # Add visit history if provided
        if visit_history:
            self._add_visit_history(doc, visit_history)
        
        # Add test results if provided
        if test_results:
            self._add_test_results(doc, test_results)
        
        # Add space for doctor's notes if requested
        if include_doctor_notes:
            self._add_doctor_notes_section(doc)
        
        # Add footer with date and page numbers
        self._add_footer(doc)
        
        # Save the document to a temporary file
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        patient_name = patient_data.get("name", "Unknown").replace(" ", "_")
        filename = f"Patient_Report_{patient_name}_{timestamp}.docx"
        file_path = os.path.join(temp_dir, filename)
        
        doc.save(file_path)
        return file_path
    
    def _setup_document(self, doc):
        """Set up document properties"""
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def _add_clinic_info(self, doc):
        """Add clinic information header"""
        # Get clinic information from config
        clinic_name = self.config_manager.config.get("clinic_name", "Medical Clinic")
        clinic_address = self.config_manager.config.get("clinic_address", "")
        clinic_phone = self.config_manager.config.get("clinic_phone", "")
        clinic_email = self.config_manager.config.get("clinic_email", "")
        
        # Create header table
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(2.0)
        
        # Add clinic information
        cell = table.cell(0, 0)
        clinic_para = cell.paragraphs[0]
        clinic_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        clinic_name_run = clinic_para.add_run(clinic_name)
        clinic_name_run.bold = True
        clinic_name_run.font.size = Pt(14)
        clinic_para.add_run("\n" + clinic_address)
        clinic_para.add_run("\nPhone: " + clinic_phone)
        clinic_para.add_run("\nEmail: " + clinic_email)
        
        # Add logo placeholder (or actual logo if available)
        cell = table.cell(0, 1)
        logo_para = cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_para.add_run("LOGO")
        
        # Add spacing after the table
        doc.add_paragraph("")
    
    def _add_title(self, doc, title):
        """Add a title to the document"""
        title_para = doc.add_heading(level=1)
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a horizontal line after the title
        self._add_horizontal_line(doc)
    
    def _add_horizontal_line(self, doc):
        """Add a horizontal line to the document"""
        p = doc.add_paragraph()
        p.add_run("_" * 80)  # Simple underscores
        p.paragraph_format.space_after = Pt(10)
        
    def _add_section_heading(self, doc, heading_text):
        """Add a section heading"""
        heading = doc.add_heading(level=2)
        heading.add_run(heading_text).bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    def _add_patient_info(self, doc, patient_data):
        """Add patient information section"""
        self._add_section_heading(doc, "Patient Information")
        
        # Create table for patient info
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'
        
        # Name
        table.cell(0, 0).text = "Name:"
        table.cell(0, 1).text = patient_data.get("name", "")
        
        # ID
        table.cell(0, 2).text = "Patient ID:"
        table.cell(0, 3).text = patient_data.get("id", "")
        
        # Date of Birth
        table.cell(1, 0).text = "Date of Birth:"
        table.cell(1, 1).text = patient_data.get("dob", "")
        
        # Gender
        table.cell(1, 2).text = "Gender:"
        table.cell(1, 3).text = patient_data.get("gender", "")
        
        # Phone
        table.cell(2, 0).text = "Phone:"
        table.cell(2, 1).text = patient_data.get("phone", "")
        
        # Email
        table.cell(2, 2).text = "Email:"
        table.cell(2, 3).text = patient_data.get("email", "")
        
        # Address
        table.cell(3, 0).text = "Address:"
        address_cell = table.cell(3, 1)
        address_cell.merge(table.cell(3, 3))
        address_cell.text = patient_data.get("address", "")
        
        # Emergency Contact
        table.cell(4, 0).text = "Emergency Contact:"
        emergency_contact = patient_data.get("emergency_contact", {})
        table.cell(4, 1).text = emergency_contact.get("name", "")
        
        table.cell(4, 2).text = "Emergency Phone:"
        table.cell(4, 3).text = emergency_contact.get("phone", "")
        
        # Insurance
        table.cell(5, 0).text = "Insurance Provider:"
        insurance = patient_data.get("insurance", {})
        table.cell(5, 1).text = insurance.get("provider", "")
        
        table.cell(5, 2).text = "Insurance ID:"
        table.cell(5, 3).text = insurance.get("id", "")
        
        # Format the table
        for row in table.rows:
            for cell in row.cells:
                # Bold the label cells
                if cell.text.endswith(":"):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        
        # Add notes if available
        patient_notes = patient_data.get("notes", "")
        if patient_notes:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.add_run("Notes: ").bold = True
            para.add_run(patient_notes)
    
    def _add_medical_history(self, doc, medical_history):
        """Add medical history section"""
        self._add_section_heading(doc, "Medical History")
        
        # Sort by date (newest first)
        sorted_history = sorted(medical_history, 
                                key=lambda x: x.get("date", ""),
                                reverse=True)
        
        # Create table
        if sorted_history:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Add headers
            headers = table.rows[0].cells
            headers[0].text = "Date"
            headers[1].text = "Type"
            headers[2].text = "Condition"
            headers[3].text = "Notes/Treatment"
            
            # Make header bold
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add medical history entries
            for entry in sorted_history:
                row = table.add_row().cells
                
                # Format date nicely
                date_str = entry.get("date", "")
                try:
                    date_obj = datetime.datetime.fromisoformat(date_str)
                    formatted_date = date_obj.strftime("%Y-%m-%d")
                except:
                    formatted_date = date_str
                
                row[0].text = formatted_date
                row[1].text = entry.get("type", "")
                row[2].text = entry.get("condition", "")
                
                # Combine notes and treatment
                notes_treatment = []
                if entry.get("notes"):
                    notes_treatment.append(entry["notes"])
                if entry.get("treatment"):
                    notes_treatment.append(f"Treatment: {entry['treatment']}")
                
                row[3].text = "\n".join(notes_treatment)
        else:
            doc.add_paragraph("No medical history data available.")
    
    def _add_visit_history(self, doc, visit_history):
        """Add visit history section"""
        self._add_section_heading(doc, "Visit History")
        
        # Sort by date (newest first)
        sorted_visits = sorted(visit_history, 
                              key=lambda x: x.get("end_time", ""),
                              reverse=True)
        
        # Create table
        if sorted_visits:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Add headers
            headers = table.rows[0].cells
            headers[0].text = "Date"
            headers[1].text = "Doctor"
            headers[2].text = "Reason"
            headers[3].text = "Diagnosis"
            headers[4].text = "Treatment"
            
            # Make header bold
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add visit entries
            for visit in sorted_visits:
                row = table.add_row().cells
                
                # Format date nicely
                date_str = visit.get("end_time", "")
                try:
                    date_obj = datetime.datetime.fromisoformat(date_str)
                    formatted_date = date_obj.strftime("%Y-%m-%d")
                except:
                    formatted_date = date_str
                
                row[0].text = formatted_date
                row[1].text = visit.get("doctor", "")
                row[2].text = visit.get("reason", "")
                row[3].text = visit.get("diagnosis", "")
                row[4].text = visit.get("treatment", "")
        else:
            doc.add_paragraph("No visit history data available.")
    
    def _add_test_results(self, doc, test_results):
        """Add test results section"""
        self._add_section_heading(doc, "Test Results")
        
        # Sort by date (newest first)
        sorted_results = sorted(test_results, 
                               key=lambda x: x.get("test_date", ""),
                               reverse=True)
        
        # Create table
        if sorted_results:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Add headers
            headers = table.rows[0].cells
            headers[0].text = "Date"
            headers[1].text = "Test Name"
            headers[2].text = "Value"
            headers[3].text = "Unit"
            headers[4].text = "Reference Range"
            headers[5].text = "Status"
            
            # Make header bold
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add test result entries
            for result in sorted_results:
                row = table.add_row().cells
                
                # Format date nicely
                date_str = result.get("test_date", "")
                try:
                    date_obj = datetime.datetime.fromisoformat(date_str)
                    formatted_date = date_obj.strftime("%Y-%m-%d")
                except:
                    formatted_date = date_str
                
                row[0].text = formatted_date
                row[1].text = result.get("test_name", "")
                row[2].text = str(result.get("value", ""))
                row[3].text = result.get("unit", "")
                row[4].text = result.get("reference_range", "")
                row[5].text = result.get("is_normal", "Unknown")
                
                # Highlight abnormal results
                if result.get("is_normal") == "Abnormal":
                    for paragraph in row[5].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        else:
            doc.add_paragraph("No test results data available.")
    
    def _add_doctor_notes_section(self, doc):
        """Add section for doctor's notes"""
        self._add_section_heading(doc, "Doctor's Notes")
        
        # Add explanatory text
        doc.add_paragraph("This section is for the doctor to fill in additional notes, observations, and recommendations.")
        
        # Add lines for writing
        for _ in range(15):
            doc.add_paragraph("_" * 90)
    
    def _add_footer(self, doc):
        """Add footer with date and page numbers"""
        # Ensure we have at least one section
        if len(doc.sections) == 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            
        section = doc.sections[0]
        footer = section.footer
        
        # Create footer paragraph
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Clear existing runs
        for run in p.runs:
            run.clear()
        
        # Add current date
        current_date = datetime.datetime.now().strftime("%Y-%m-%d")
        p.add_run(f"Generated on: {current_date} | ")
        
        # Add page numbers
        p.add_run("Page ")
        
        # Add a special XML element for page number
        run = p.add_run()
        r = run._r
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        instrText.set(qn('xml:space'), 'preserve')
        r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r.append(fldChar2)
        
        # Total pages
        p.add_run(" of ")
        
        # Total pages XML element
        run = p.add_run()
        r = run._r
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'NUMPAGES'
        instrText.set(qn('xml:space'), 'preserve')
        r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r.append(fldChar2)


    def _add_medications(self, doc, medications):
        """Add medications section to the report"""
        self._add_section_heading(doc, "Current Medications")
        
        # Sort by name
        sorted_meds = sorted(medications, key=lambda x: x.get("name", ""))
        
        # Create table
        if sorted_meds:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Add headers
            headers = table.rows[0].cells
            headers[0].text = "Medication"
            headers[1].text = "Dosage"
            headers[2].text = "Frequency"
            headers[3].text = "Start Date"
            headers[4].text = "Notes"
            
            # Make header bold
            for cell in headers:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add medication entries
            for med in sorted_meds:
                row = table.add_row().cells
                
                row[0].text = med.get("name", "")
                row[1].text = med.get("dosage", "")
                row[2].text = med.get("frequency", "")
                row[3].text = med.get("start_date", "")
                
                # Combine notes with end date info if present
                notes = med.get("notes", "")
                end_date = med.get("end_date", "")
                
                if end_date:
                    if notes:
                        notes = f"End date: {end_date}\n\n{notes}"
                    else:
                        notes = f"End date: {end_date}"
                        
                row[4].text = notes
        else:
            doc.add_paragraph("No medication information available.")